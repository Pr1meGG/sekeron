import re
from pathlib import Path
from typing import Tuple, List, Optional, Dict, Any
import docx
from .models import ProfileClaim, DataQualityIssue

def parse_docx_text(path: Path) -> str:
    """Read a docx file and return its raw text content."""
    try:
        doc = docx.Document(str(path))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())
        # Also parse tables if text is within them
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in paragraphs:
                        paragraphs.append(cell_text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to read docx file {path.name}: {str(e)}")

def extract_profile_details(text: str, folder_name: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str], Optional[str], str]:
    """
    Extract ID, Name, Category, Location, Work Preference, and Bio from profile text.
    Uses folder name as a fallback for missing display names.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    claimed_id = None
    claimed_name = None
    category = None
    location = None
    work_pref = None
    bio = ""
    
    # 1. ID Extraction using regex
    id_match = re.search(r'\b([PVM]O?\d+)\b', text)
    if id_match:
        claimed_id = id_match.group(1)
        
    # 2. Display Name Extraction
    # Case: "Artist Name -Jane Doe"
    name_match = re.search(r'Artist Name\s*[-—:]\s*([A-Za-z\s\'&]+)', text, re.IGNORECASE)
    if name_match:
        claimed_name = name_match.group(1).strip()
    else:
        # Check first line for common header formats:
        if lines and not lines[0].startswith("Category:") and not lines[0].startswith("Location:"):
            first_line = lines[0]
            # Case "P01 / Jane Doe" or "M01 — John & Jane" or "V03 / Jane Smith"
            if '/' in first_line:
                parts = first_line.split('/')
                claimed_name = parts[1].strip()
            elif '—' in first_line:
                parts = first_line.split('—')
                claimed_name = parts[1].strip()
            elif '-' in first_line:
                parts = first_line.split('-')
                claimed_name = parts[1].strip()
            elif '(' in first_line and ')' in first_line: # Case "Jane Doe (M03)"
                claimed_name = re.sub(r'\([A-Za-z0-9\s]+\)', '', first_line).strip()
            elif claimed_id and claimed_id in first_line:
                # E.g. "M04_UniqueName" -> split and clean
                cleaned = first_line.replace(claimed_id, "").replace("_", " ").strip()
                claimed_name = cleaned
            else:
                # Fallback to the first line itself if it's not too long and doesn't match ID patterns
                if len(first_line) < 50:
                    claimed_name = first_line.strip()
                    
    # Clean name
    if claimed_name:
        # Remove trailing info like "Category" or "Artist ID"
        claimed_name = re.sub(r'\b(Category|Location|Artist ID|Artist Name|Acoustic duo|Bio)\b.*', '', claimed_name, flags=re.IGNORECASE).strip()
        
    # 3. Category, Location, Work Preference extraction
    # Standard format: "Category: Photographer Location: Delhi / NCR Work preference: Travel available"
    cat_match = re.search(r'Category\s*[-—:]?\s*([A-Za-z\s/&]+?)(?=\bLocation:|\bWork preference:|\n|$)', text, re.IGNORECASE)
    if cat_match:
        category = cat_match.group(1).strip()
        
    loc_match = re.search(r'Location\s*[-—:]?\s*([A-Za-z\s/,\.\(\)]+?)(?=\bWork preference:|\bCategory:|\n|$)', text, re.IGNORECASE)
    if loc_match:
        location = loc_match.group(1).strip()
        
    pref_match = re.search(r'Work preference\s*[-—:]?\s*([A-Za-z\s/,\.&]+?)(?=\n|Bio:|$)', text, re.IGNORECASE)
    if pref_match:
        work_pref = pref_match.group(1).strip()
        
    # 4. Bio extraction
    bio_match = re.search(r'Bio\s*[-—:]\s*(.+?)(?=\bPortfolio:|\bWork preference:|\n\n|$)', text, re.DOTALL | re.IGNORECASE)
    if bio_match:
        bio = bio_match.group(1).strip()
    else:
        # Find any text after "Bio:" or fallback to lines containing biography sentences
        for line in lines:
            if line.startswith("Bio:"):
                bio = line.replace("Bio:", "").strip()
                break
        if not bio and len(lines) > 2:
            # Simple heuristic: last few lines that are long
            bio = lines[-1]
            
    return claimed_id, claimed_name, category, location, work_pref, bio

def extract_claims_and_issues(
    docx_path: Path, 
    artist_key: str, 
    folder_name: str
) -> Tuple[List[ProfileClaim], List[DataQualityIssue], Dict[str, Any]]:
    """
    Parse a docx profile, extract capabilities and demographic profile claims,
    and detect configuration/identity conflicts.
    """
    text = parse_docx_text(docx_path)
    claimed_id, claimed_name, category, location, work_pref, bio = extract_profile_details(text, folder_name)
    
    claims = []
    issues = []
    
    source_file = str(docx_path)
    
    # Generate structured Profile Claims
    if category:
        claims.append(ProfileClaim(
            artist_key=artist_key,
            claim=category,
            source_file=source_file,
            section_or_locator="Profile Header (Category)",
            claim_type="category",
            confidence=1.0,
            notes=f"Claimed category in profile text"
        ))
    if location:
        claims.append(ProfileClaim(
            artist_key=artist_key,
            claim=location,
            source_file=source_file,
            section_or_locator="Profile Header (Location)",
            claim_type="location",
            confidence=1.0,
            notes=f"Claimed location in profile text"
        ))
    if work_pref:
        claims.append(ProfileClaim(
            artist_key=artist_key,
            claim=work_pref,
            source_file=source_file,
            section_or_locator="Profile Header (Work Preference)",
            claim_type="work_preference",
            confidence=1.0,
            notes=f"Claimed work preference in profile text"
        ))
        
    if bio:
        claims.append(ProfileClaim(
            artist_key=artist_key,
            claim=bio,
            source_file=source_file,
            section_or_locator="Profile Body (Bio)",
            claim_type="biography",
            confidence=1.0,
            notes="Full profile biography statement"
        ))
        
        # Parse explicit capabilities or skills mentioned in bio
        # Split on common delimiters or scan for comma lists
        # E.g. "working in events, cafés, workshops..."
        skills_match = re.search(r'(?:working in|creating|working on)\s+([^.]+)', bio, re.IGNORECASE)
        if skills_match:
            skills_text = skills_match.group(1)
            # split by commas, 'and', or semicolons
            skills = [s.strip() for s in re.split(r',|and|;', skills_text) if s.strip()]
            for skill in skills:
                claims.append(ProfileClaim(
                    artist_key=artist_key,
                    claim=skill,
                    source_file=source_file,
                    section_or_locator="Profile Body (Bio Skills List)",
                    claim_type="capability_claim",
                    confidence=0.8,
                    notes=f"Explicit capability claim parsed from biography"
                ))

    # Identify identity and naming anomalies
    # Anomaly 1: Missing Name
    if not claimed_name:
        derived_name = folder_name.split('_', 1)[1].replace('_', ' ') if '_' in folder_name else folder_name
        issues.append(DataQualityIssue(
            severity="WARNING",
            issue_type="missing_profile_name",
            artist_key=artist_key,
            asset_key=None,
            description="The profile document is missing a name in the header.",
            evidence=f"Profile text: {text[:100]}..."
        ))
        display_name = derived_name
    else:
        display_name = claimed_name

    # Anomaly 2: Mismatched folder ID vs claimed ID
    folder_id = folder_name.split('_')[0]
    # normalize O/0
    norm_folder_id = folder_id.replace('O', '0')
    if claimed_id:
        norm_claimed_id = claimed_id.replace('O', '0')
        if norm_folder_id != norm_claimed_id:
            issues.append(DataQualityIssue(
                severity="ERROR",
                issue_type="id_conflict",
                artist_key=artist_key,
                asset_key=None,
                description=f"Folder ID '{folder_id}' does not match Claimed ID '{claimed_id}' in profile.",
                evidence=f"Folder: {folder_name}, Profile claim: ID {claimed_id}"
            ))

    # Anomaly 3: Mismatched folder name vs claimed name
    folder_name_core = folder_name.split('_', 1)[1] if '_' in folder_name else folder_name
    norm_folder_name = folder_name_core.replace('_', ' ').lower().strip()
    if claimed_name:
        norm_claimed_name = claimed_name.lower().strip()
        # check if they are significantly different (e.g. Jane Smith vs John Doe)
        if norm_claimed_name not in norm_folder_name and norm_folder_name not in norm_claimed_name:
            issues.append(DataQualityIssue(
                severity="ERROR",
                issue_type="name_mismatch",
                artist_key=artist_key,
                asset_key=None,
                description=f"Folder name '{folder_name_core}' does not match Claimed name '{claimed_name}' in profile.",
                evidence=f"Folder: {folder_name}, Profile claim: Name {claimed_name}"
            ))

    # Collect details for return
    profile_details = {
        "claimed_id": claimed_id,
        "display_name": display_name,
        "claimed_category": category,
        "claimed_location": location,
        "claimed_work_preference": work_pref,
        "bio": bio,
        "full_text": text
    }
    
    return claims, issues, profile_details
